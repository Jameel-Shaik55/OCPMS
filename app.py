from flask import Flask, render_template, request, session, redirect, url_for, flash, send_file, jsonify
import pandas as pd
from flask_mail import *
import secrets
import os
from datetime import datetime
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import secrets
from werkzeug.utils import secure_filename
import mysql.connector
from datetime import datetime
import os
from flask import render_template
from werkzeug.utils import secure_filename
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from reportlab.lib.pagesizes import letter
from datetime import timedelta


mydb = mysql.connector.connect(
    host="localhost", user="root", passwd="admin@123", database="campus", charset='utf8', port=3306)
mycursor = mydb.cursor()

UPLOAD_FOLDER = '/path/to/the/uploads'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif'}

app = Flask(__name__)
app.secret_key = 'your_secret_key_here'
app.config['UPLOAD_FOLDER'] = 'static/uploads'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'doc', 'docx', 'jpg', 'jpeg', 'png'}
app.permanent_session_lifetime = timedelta(hours=2)



# home page
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/signin', methods=['GET', 'POST'])
def signin():
    if request.method == "POST":
        user_type = request.form['user_type']
        email = request.form['email']
        password = request.form['password']
        print(password)

        if user_type == 'company':
            sql = "SELECT * FROM company WHERE email=%s AND password=%s"
            mycursor.execute(sql, (email, password))
            results = mycursor.fetchall()

            if len(results) > 0:
                company_name = results[0][1]
                session['compyemail'] = email
                session['company_name'] = company_name
                return render_template('companyhome.html', msg="Login successful")
            else:
                return render_template('index.html', msg="Login failed")

        elif user_type == 'tpo':
            if email == 'tpo@gmail.com' and password == 'tpo':
                return render_template('tpohome.html', msg="Login successful")
            else:
                return render_template('index.html', msg="Login failed")

        elif user_type == 'student':
            session['student'] = email
            session.permanent = True
            sql = "SELECT * FROM students WHERE email=%s AND password=%s"
            mycursor.execute(sql, (email, password))
            results = mycursor.fetchall()

            if len(results) > 0:
                session['studentsemail'] = email
                return render_template('studenthome.html', msg="Login successful")
            else:
                return render_template('index.html', msg="Login failed")
    
    return render_template('index.html')


@app.route("/forgotpassword", methods=['POST', 'GET'])
def forgotpassword():
    if request.method == "POST":
        user_type = request.form['user_type']
        email = request.form['email']
        if user_type == 'comapny':
            sql = "SELECT * FROM company WHERE email=%s"
            mycursor.execute(sql, (email,))
            data = mycursor.fetchall()
            mydb.commit()
            if data:
                msg = 'valid'
                session['sforgotemail'] = email
                return render_template('forgotpassword.html', msg=msg)
            else:
                msg = "notvalid"
                flash("Provide Valid Email", "warning")
                return render_template('index.html', msg=msg)
        elif user_type == 'student':
            sql = "SELECT * FROM students WHERE email=%s"
            mycursor.execute(sql, (email,))
            data = mycursor.fetchall()
            mydb.commit()
            if data:
                msg = 'valid'
                session['bforgotemail'] = email
                return render_template('bforgotpassword.html', msg=msg)
            else:
                msg = "notvalid"
                flash("Provide Valid Email", "warning")
                return render_template('index.html', msg=msg)
    return render_template('forgotpassword.html', msg='check')

@app.route("/updatepassword", methods=['POST'])
def updatepassword():
    if request.method == "POST":
        email = session.get('sforgotemail') or session.get('bforgotemail')
        password = request.form['password']
        confirmpassword = request.form['confirmpassword']

        if password == confirmpassword:
            user_type = 'comapny' if 'sforgotemail' in session else 'student'
            email_field = 'email' if user_type == 'student' else 'bemail'
            sql = f"SELECT * FROM {user_type} WHERE {email_field}=%s"
            mycursor.execute(sql, (email,))
            data = mycursor.fetchall()
            mydb.commit()
            if data:
                sql = f"UPDATE {user_type} SET password=%s WHERE {email_field}=%s"
                mycursor.execute(sql, (password, email))
                mydb.commit()
                flash("Password Updated Successfully", "success")
                return redirect(url_for("/" if user_type == 'sellers' else "/"))
        else:
            flash("Passwords do not match", "warning")
            return render_template("index.html" if 'sforgotemail' in session else "buyerslog.html")
    return render_template("index.html" if 'sforgotemail' in session else "buyerslog.html")

@app.route('/companyregister', methods=['GET', 'POST'])
def companyregister():
    if request.method == "POST":
        company_name = request.form['company_name']
        email = request.form['email']
        password = request.form['password']
        con_password = request.form['con_password']
        phone = request.form['phone']
        address = request.form['address']
        city = request.form['city']
        state = request.form['state']
        postal_code = request.form['postal_code']
        country = request.form['country']
        website = request.form.get('website', '')  # Optional field
        description = request.form['description']
        logo = request.files.get('logo')

        # Check if passwords match
        if password != con_password:
            flash('Passwords do not match', 'danger')
            return redirect(url_for('companyregister'))

        # Check if company name or email already exists
        mycursor.execute("SELECT * FROM company WHERE company_name = %s OR email = %s", (company_name, email))
        existing_record = mycursor.fetchone()
        if existing_record:
            flash('Company name or email already exists, please use a different one.', 'danger')
            return redirect(url_for('companyregister'))

        # Save the logo file if provided
        logo_filename = None
        if logo:
            logo_filename = secure_filename(logo.filename)
            logo.save(os.path.join('static/uploads', logo_filename))

        # Insert new company record
        sql = """
        INSERT INTO company (company_name, email,password, phone, address, city, state, postal_code, country, website, description, logo)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """
        values = (company_name, email,password, phone, address, city, state, postal_code, country, website, description, logo_filename)
        
        try:
            mycursor.execute(sql, values)
            mydb.commit()
            flash('Company registration successful', 'success')
            return redirect(url_for('signin'))
        except Exception as e:
            mydb.rollback()
            flash(f'Error: {str(e)}', 'danger')
    return render_template('companyregister.html')





# Add  students
@app.route("/student", methods=["POST", "GET"])
def student():
    profilepath = ""  # Initialize profilepath with a default value
    if request.method == "POST":
        name = request.form['name']
        email = request.form['email']
        password = request.form['password']
        password1 = request.form['Con_Password']
        rollno = request.form['rollno']
        year = request.form['year']
        sem = request.form['sem']
        branch = request.form['branch']
        contact = request.form['mobile']
        address = request.form['address']
        myfile = request.files['myfile']
        filename = myfile.filename  # Define filename here
        path = os.path.join("static/profiles/", filename)
        myfile.save(path)
        profilepath = "static/profiles/" + filename
        now = datetime.now()
        current_datetime = now.strftime("%Y-%m-%d %H:%M:%S")
        if password == password1:
            sql = "select * from students where email='%s' and password='%s'" % (
                email, password)
            mycursor.execute(sql)
            data = mycursor.fetchall()
            print(data)
            if data == []:
                print(name, email, password, address)
                sql = "insert into students(name,email,rollno,year,sem,branch,password,contact,address,Datetime,profile) values(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
                val = (name, email, rollno, year, sem, branch, password, contact,
                        address,  current_datetime, profilepath)
                mycursor.execute(sql, val)
                mydb.commit()
                return redirect(url_for('signin'))
            else:
                flash('Details already Exist', "warning")
                return render_template('student.html', msg="Details already Exist")
        else:
            flash('password not matched')
            return render_template('student.html')
    return render_template('student.html')


@app.route('/add_question', methods=['GET', 'POST'])
def add_question(): 
    
    return render_template('add_question.html')


@app.route('/add_mcq', methods=['GET', 'POST'])
def add_mcq():
    if 'compyemail' not in session:
        return redirect(url_for('companysignin'))

    if request.method == "POST":
        question = request.form['question']
        option1 = request.form['option1']
        option2 = request.form['option2']
        option3 = request.form['option3']
        option4 = request.form['option4']
        correct_option = request.form['correct_option']

        # Insert question into database
        sql = """INSERT INTO mcqquestion (company_email, question_text, option1, option2, option3, option4, correct_option, stage) 
                 VALUES (%s, %s, %s, %s, %s, %s, %s, 'MCQ')"""
        mycursor.execute(sql, (session['compyemail'], question, option1, option2, option3, option4, correct_option))
        mydb.commit()
        
        flash("Question added successfully!", "success")
        return redirect(url_for('add_mcq'))
    
    return render_template('add_mcq.html')

@app.route('/addcodingtest',  methods=['GET', 'POST'])
def addcodingtest():
    if 'compyemail' not in session:
        return redirect(url_for('companysignin'))
    if request.method == "POST":
        question = request.form['coding_question']
        sample_input = request.form['sample_input']
        expected_output = request.form['expected_output']
        language = request.form['language']
    
        sql = "INSERT INTO coding_tests (company_email, question_text, sample_input, expected_output,language) VALUES (%s, %s, %s, %s, %s)"
        mycursor.execute(sql, (session['compyemail'], question, sample_input, expected_output,language))
        mydb.commit()
    
        flash('Coding test question added successfully!', 'success')
   
    return render_template('add_coding.html')


@app.route('/manage_question', methods=['GET', 'POST'])
def manage_question():
    return render_template('manage_question.html')

# Route to view MCQ questions
@app.route('/view_mcq_questions')
def view_mcq_questions():
    sql = "SELECT * FROM mcqquestion WHERE company_email = %s"
    company_email = session['compyemail']  # Temporarily hardcoded for testing purposes
    mycursor.execute(sql, (company_email,))
    data = mycursor.fetchall()
    print(data)  # Debugging line to check data fetched
    return render_template('view_mcq_questions.html', data=data)

@app.route('/mcqdelete/<Id>')
def mcqdelete(Id=0):
    sql = "delete from mcqquestion where Id='%s' " % (Id)
    mycursor.execute(sql)
    mydb.commit()
    return redirect(url_for('view_mcq_questions'))

@app.route("/updatemcq/<int:Id>")
def updatemcq(Id=0):
    sql = "select * from mcqquestion where Id='%s'"%(Id)
    mycursor.execute(sql)
    data = mycursor.fetchall()
    print(data)
    return render_template('updatemcq.html',Id=data[0][0],data=data)

@app.route('/updatemcqback', methods=["POST", "GET"])
def updatemcqback():
    if request.method == "POST":
        # Retrieve form data
        Id = request.form['Id']
        question = request.form['question']
        option1 = request.form['option1']
        option2 = request.form['option2']
        option3 = request.form['option3']
        option4 = request.form['option4']
        correct_option = request.form['correct_option']
        
        # SQL update statement with WHERE clause
        sql = """UPDATE mcqquestion 
                 SET question_text=%s, option1=%s, option2=%s, option3=%s, option4=%s, correct_option=%s 
                 WHERE Id=%s"""
        val = (question, option1, option2, option3, option4, correct_option, Id)

        # Execute the update statement
        mycursor.execute(sql, val)
        mydb.commit()

        # Redirect to the viewcropinfo page after updating
        return redirect(url_for('view_mcq_questions'))

    # Handle GET request or other methods if needed
    return redirect(url_for('view_mcq_questions'))

@app.route('/view_code_questions')
def view_code_questions():
    sql = "SELECT * FROM coding_tests WHERE company_email = %s"
    company_email = session['compyemail']  # Temporarily hardcoded for testing purposes
    mycursor.execute(sql, (company_email,))
    data = mycursor.fetchall()
    print(data)  # Debugging line to check data fetched
    return render_template('view_code_questions.html', data=data)

@app.route('/codedelete/<Id>')
def codedelete(Id=0):
    sql = "delete from coding_tests where Id='%s' " % (Id)
    mycursor.execute(sql)
    mydb.commit()
    return redirect(url_for('view_code_questions'))

@app.route("/updatecode/<int:Id>")
def updatecode(Id=0):
    sql = "SELECT * FROM coding_tests WHERE Id=%s"
    mycursor.execute(sql, (Id,))
    data = mycursor.fetchall()
    print(data)
    return render_template('updatecode.html', Id=data[0][0], data=data)

@app.route('/updatecodeback', methods=["POST", "GET"])
def updatecodeback():
    if request.method == "POST":
        # Retrieve form data
        Id = request.form['Id']
        question = request.form['coding_question']
        sample_input = request.form['sample_input']
        expected_output = request.form['expected_output']
        language = request.form['language']
        # SQL update statement with WHERE clause
        sql = """UPDATE coding_tests 
                 SET question_text=%s, sample_input=%s, expected_output=%s, language=%s 
                 WHERE Id=%s"""
        val = (question, sample_input, expected_output, language, Id)

        # Execute the update statement
        mycursor.execute(sql, val)
        mydb.commit()

        # Redirect to the view_code_questions page after updating
        return redirect(url_for('view_code_questions'))

    # Handle GET request or other methods if needed
    return redirect(url_for('view_code_questions'))

    

@app.route('/add_drive', methods=['GET', 'POST'])
def add_drive():
    sql = "select * from company where email='%s'"%(session['compyemail'])
    mycursor.execute(sql)
    dc = mycursor.fetchall()
    print(dc)
    company_name = dc[0][1]
    website = dc[0][10]
    logo = dc[0][12]
    print(company_name,website,logo)
    
    if 'compyemail' not in session:
        return redirect(url_for('companysignin'))
    
    if request.method == "POST":
        date_of_conduct = request.form['date_of_conduct']
        position = request.form['position']
        location = request.form['location']
        job_vacancies = request.form['job_vacancies']
        job_description = request.form['job_description']
        salary = request.form['salary']
        department = request.form['department']
        
        # Check if the same drive already exists within the last 6 months
        sql_check = """SELECT * FROM drive WHERE company_email = %s AND position = %s 
                       AND date_of_conduct >= DATE_SUB(%s, INTERVAL 6 MONTH)"""
        mycursor.execute(sql_check, (session['compyemail'], position, date_of_conduct))
        existing_drive = mycursor.fetchone()
        
        if existing_drive:
            flash("A similar drive has already been added in the last 6 months!", "danger")
        else:
            # Insert drive into database
            sql = """INSERT INTO drive (company_email, date_of_conduct, position, location, job_vacancies, job_description, salary, department,company_name,website,logo) 
                     VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)"""
            mycursor.execute(sql, (session['compyemail'], date_of_conduct, position, location, job_vacancies, job_description, salary, department,company_name,website,logo))
            mydb.commit()
            flash("Drive added successfully!", "success")
            
        return render_template('add_drive.html')
    return render_template('add_drive.html')


@app.route("/viewdrive")
def viewdrive():
    if 'compyemail' not in session:
        return redirect(url_for('companysignin'))

    # Fetch drive data
    sql_drive = "SELECT * FROM drive WHERE company_email=%s"
    data = pd.read_sql_query(sql_drive, mydb, params=(session['compyemail'],))

    return render_template('viewdrive.html', cols=data.columns.values, rows=data.values.tolist())

@app.route("/viewalldrive")
def viewalldrive():
    sql_drive = "SELECT * FROM drive"
    data = pd.read_sql_query(sql_drive, mydb)
    return render_template('viewalldrive.html', cols=data.columns.values, rows=data.values.tolist())



@app.route('/driverequest/<int:id>', methods=["POST", "GET"])
def driverequest(id):
    # Fetch drive details from the database
    sql = "SELECT * FROM drive WHERE id=%s"
    mycursor.execute(sql, (id,))
    data = mycursor.fetchone()

    if data:
        company_email = data[1]
        date_of_conduct = data[2]
        position = data[3]
        location = data[4]
        job_vacancies = data[5]
        job_description = data[6]
        salary = data[7]
        department = data[8]
        company_name = data[9]
        website = data[10]
        logo = data[11]

        # Check for duplicate requests
        sql_check = "SELECT * FROM driverequest WHERE company_email=%s AND position=%s"
        mycursor.execute(sql_check, (company_email, position,))
        existing_request = mycursor.fetchone()

        if existing_request:
            flash("A request for this position has already been sent.")
        else:
            # Insert new request into driverequest table
            sql_insert = """
                INSERT INTO driverequest (company_name, company_email, date_of_conduct, position, location, 
                job_vacancies, job_description, salary, department, website, logo)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """
            val_insert = (company_name, company_email, date_of_conduct, position, location, 
                          job_vacancies, job_description, salary, department, website, logo)
            mycursor.execute(sql_insert, val_insert)
            mydb.commit()
            flash("Request sent successfully.")

    return redirect(url_for('viewalldrive'))

@app.route("/viewrequestback")
def viewrequestback():
    if 'company_name' not in session:
        return redirect(url_for('companysignin'))

   

    return render_template('viewrequestback.html')

@app.route("/viewrequest")
def viewrequest():
    if 'company_name' not in session:
        return redirect(url_for('companysignin'))

    # Fetch drive data
    sql_drive = "SELECT * FROM driverequest WHERE company_name=%s"
    data = pd.read_sql_query(sql_drive, mydb, params=(session['company_name'],))

    return render_template('viewrequest.html', cols=data.columns.values, rows=data.values.tolist())

@app.route("/accept/<id>")
def accept(id=0):
    print(id)
    sql = "select  * from driverequest where id='%s'" % (id)
    mycursor.execute(sql)
    dc = mycursor.fetchall()
    print(dc)
    collegename = dc[0][1]
    collegeemail = dc[0][2]
    print(collegename, collegeemail)
    skey = secrets.token_hex(4)
    print("secret key", skey)
    mail_content = 'Your request for the drive for ' + collegename + ' College is accepted by our company and email is: ' + collegeemail + ' '
    sender_address = 'cse.takeoff@gmail.com'
    sender_pass = 'digkagfgyxcjltup'
    receiver_address = collegeemail
    message = MIMEMultipart()
    message['From'] = sender_address
    message['To'] = receiver_address
    message['Subject'] = 'Campus Placement Recruitment System'
    message.attach(MIMEText(mail_content, 'plain'))
    session_email = smtplib.SMTP('smtp.gmail.com', 587)
    session_email.starttls()
    session_email.login(sender_address, sender_pass)
    text = message.as_string()
    session_email.sendmail(sender_address, receiver_address, text)
    session_email.quit()
    sql = "update driverequest set Status='Approved' where id='%s'" % (id)
    mycursor.execute(sql)
    mydb.commit()
    flash("Drive request accepted successfully.", "success")
    return redirect(url_for('viewrequest'))

@app.route("/rejected/<id>")
def rejected(id=0):
    print(id)
    sql = "select  * from driverequest where id='%s'" % (id)
    mycursor.execute(sql)
    dc = mycursor.fetchall()
    print(dc)
    collegename = dc[0][1]
    collegeemail = dc[0][2]
    print(collegename, collegeemail)
    otp = "Your Request is Rejected and this is your secret Key :"
    skey = secrets.token_hex(4)
    print("secret key", skey)
    mail_content = 'Your request for the drive for ' + collegename + ' College is rejected by our company. Please try again.'
    sender_address = 'cse.takeoff@gmail.com'
    sender_pass = 'digkagfgyxcjltup'
    receiver_address = collegeemail
    message = MIMEMultipart()
    message['From'] = sender_address
    message['To'] = receiver_address
    message['Subject'] = 'Campus Placement Recruitment System'
    message.attach(MIMEText(mail_content, 'plain'))
    session_email = smtplib.SMTP('smtp.gmail.com', 587)
    session_email.starttls()
    session_email.login(sender_address, sender_pass)
    text = message.as_string()
    session_email.sendmail(sender_address, receiver_address, text)
    session_email.quit()
    sql = "update driverequest set Status='Rejected' where id='%s'" % (id)
    mycursor.execute(sql)
    mydb.commit()
    flash("Drive request rejected successfully.", "danger")
    return redirect(url_for('viewrequest'))

@app.route("/viewrequeststatus")
def viewrequeststatus():
    # Fetch drive data
    sql_drive = "SELECT * FROM driverequest WHERE Status='Approved'"
    data = pd.read_sql_query(sql_drive, mydb)
    print(data)
    return render_template('viewrequeststatus.html', cols=data.columns.values, rows=data.values.tolist())

@app.route("/viewdrivedetails")
def viewdrivedetails():
    # Fetch drive data
    sql_drive = "SELECT * FROM driverequest WHERE Status='Approved'"
    data = pd.read_sql_query(sql_drive, mydb)
    print(data)
    return render_template('viewdrivedetails.html', cols=data.columns.values, rows=data.values.tolist())


@app.route('/applydrive<id>', methods=['GET', 'POST'])
def applydrive(id=0):
    # Fetch drive request details
    sql = "SELECT * FROM driverequest WHERE id=%s" % (id)
    mycursor.execute(sql)
    data = mycursor.fetchall()

    company_name = data[0][1]
    company_email = data[0][2]
    date_of_conduct = data[0][3]
    job_vacancies = data[0][6]
    job_description = data[0][7]
    salary = data[0][8]
    department = data[0][9]
    website = data[0][10]
    logo = data[0][11]

    # Fetch student details
    sql1 = "SELECT * FROM students WHERE email=%s"
    mycursor.execute(sql1, (session['studentsemail'],))
    dc = mycursor.fetchall()

    Id = dc[0][0]
    stdname = dc[0][1]
    stdemail = dc[0][2]  # Corrected initialization for student email
    rollno = dc[0][3]
    branch = dc[0][6]
    profile = dc[0][11]
    year = dc[0][4]

    # Fetch positions from the database
    sql_positions = "SELECT DISTINCT position FROM driverequest"
    mycursor.execute(sql_positions)
    positions = [pos[0] for pos in mycursor.fetchall()]

    # Predefined list of Indian locations
    locations = [
        "New Delhi", "Mumbai", "Bangalore", "Hyderabad", "Chennai", 
        "Kolkata", "Pune", "Ahmedabad", "Jaipur", "Lucknow", 
        "Kanpur", "Nagpur", "Visakhapatnam", "Indore", "Patna",
        "Bhopal", "Thane", "Vadodara", "Ghaziabad", "Ludhiana",
        "Agra", "Nashik", "Faridabad", "Meerut", "Rajkot",
        "Varanasi", "Srinagar", "Aurangabad", "Dhanbad", "Amritsar",
        "Navi Mumbai", "Allahabad", "Ranchi", "Howrah", "Coimbatore",
        "Jabalpur", "Gwalior", "Vijayawada", "Jodhpur", "Madurai",
        "Raipur", "Kota", "Chandigarh", "Guwahati", "Solapur"
    ]

    if request.method == "POST":
        # Handle form submission
        stdname = request.form['stdname']
        stdemail = request.form['stdemail']  # Corrected typo
        year = request.form['year']
        tenth_marks = request.form['tenth_marks']
        twelfth_marks = request.form['twelfth_marks']
        current_marks = request.form['current_marks']
        collegename = request.form['collegename']
        department = request.form['department']
        position = request.form['position']  # Dropdown value
        location = request.form['location']  # Dropdown value
        resume = request.files['resume']
        filename = resume.filename
        path = os.path.join("static/profiles/", filename)
        resume.save(path)
        profilepath = "static/profiles/" + filename
        Status = "Applied"

        # Check if the same drive already exists within the last 6 months
        sql_check = """SELECT * FROM studentrequest WHERE company_email = %s AND position = %s AND student_email = %s
                       AND date_of_conduct >= DATE_SUB(%s, INTERVAL 6 MONTH)"""
        mycursor.execute(sql_check, (company_email, position, session['studentsemail'], date_of_conduct))

        existing_drive = mycursor.fetchone()

        if existing_drive:
            flash("A similar drive has already been added in the last 6 months!", "danger")
        else:
            # Insert drive into database
            sql = """INSERT INTO studentrequest (stdname, student_email, year, tenth_marks, twelfth_marks,current_marks, collegename, department, 
                                                 position, location, company_email, date_of_conduct, 
                                                 profile, resume, Status) 
                     VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s ,%s)"""
            mycursor.execute(sql, (stdname, stdemail, year, tenth_marks, twelfth_marks, current_marks, collegename, department, 
                                   position, location, company_email, date_of_conduct, profile, profilepath, Status))
            mydb.commit()
            flash("Drive added successfully!", "success")
            
        return render_template('viewdrivedetails.html')
    
    # Pass all required variables to the template
    return render_template('applydrive.html', stdname=stdname, stdemail=stdemail, year=year, positions=positions, locations=locations)

@app.route("/viewrequeststd")
def viewrequeststd():
    # Fetch drive data
    sql_drive = "SELECT * FROM studentrequest WHERE company_email=%s "
    data = pd.read_sql_query(sql_drive, mydb, params=(session.get('compyemail'),))
    
    # Debug output
    print("Fetched data:")
    print(data.head())  # Print the first few rows to verify the paths
    
    return render_template('viewrequeststd.html', cols=data.columns.values, rows=data.values.tolist())

@app.route("/stdaccept/<id>")
def stdaccept(id=0):
    print(id)
    sql = "select  * from studentrequest where id='%s'" % (id)
    mycursor.execute(sql)
    dc = mycursor.fetchall()
    print(dc)
    stdname = dc[0][1]
    stdemail = dc[0][2]
    print(stdname, stdemail)
    skey = secrets.token_hex(4)
    print("secret key", skey)
    mail_content = 'Your request for the drive for ' + stdname + ' Student  is accepted by our company and And Procied for the exam email is: ' + stdemail + ' '
    sender_address = 'cse.takeoff@gmail.com'
    sender_pass = 'digkagfgyxcjltup'
    receiver_address = stdemail
    message = MIMEMultipart()
    message['From'] = sender_address
    message['To'] = receiver_address
    message['Subject'] = 'Campus Placement Recruitment System'
    message.attach(MIMEText(mail_content, 'plain'))
    session_email = smtplib.SMTP('smtp.gmail.com', 587)
    session_email.starttls()
    session_email.login(sender_address, sender_pass)
    text = message.as_string()
    session_email.sendmail(sender_address, receiver_address, text)
    session_email.quit()
    sql = "update studentrequest set Status='Accepted' where id='%s'" % (id)
    mycursor.execute(sql)
    mydb.commit()
    flash("Drive request accepted successfully.", "success")
    return redirect(url_for('viewrequeststd'))

@app.route("/stdrejected/<id>")
def stdrejected(id=0):
    print(id)
    sql = "select  * from studentrequest where id='%s'" % (id)
    mycursor.execute(sql)
    dc = mycursor.fetchall()
    print(dc)
    stdname = dc[0][1]
    stdemail = dc[0][2]
    print(stdname, stdemail)
    otp = "Your Request is Rejected and this is your secret Key :"
    skey = secrets.token_hex(4)
    print("secret key", skey)
    mail_content = 'Your request for the drive for ' + stdname + ' College is rejected by our company. Please try again.'
    sender_address = 'cse.takeoff@gmail.com'
    sender_pass = 'digkagfgyxcjltup'
    receiver_address = stdemail
    message = MIMEMultipart()
    message['From'] = sender_address
    message['To'] = receiver_address
    message['Subject'] = 'Campus Placement Recruitment System'
    message.attach(MIMEText(mail_content, 'plain'))
    session_email = smtplib.SMTP('smtp.gmail.com', 587)
    session_email.starttls()
    session_email.login(sender_address, sender_pass)
    text = message.as_string()
    session_email.sendmail(sender_address, receiver_address, text)
    session_email.quit()
    sql = "update studentrequest set Status='Rejected' where id='%s'" % (id)
    mycursor.execute(sql)
    mydb.commit()
    flash("Drive request rejected successfully.", "danger")
    return redirect(url_for('viewrequeststd'))


@app.route("/viewmcqquestions")
def viewmcqquestions():
    # Check if company email is in the session
    if 'compyemail' not in session:
        return redirect(url_for('companyregister'))

    # Fetch accepted student requests
    sql_student_requests = """
        SELECT * FROM studentrequest 
        WHERE company_email=%s AND Status='Accepted'
    """
    student_requests = pd.read_sql_query(sql_student_requests, mydb, params=(session['compyemail'],))
    # cursor = mydb.cursor(dictionary=True)
    # cursor.execute(sql_student_requests, (session['compyemail'],))
    # rows = cursor.fetchall()
    # student_requests = pd.DataFrame(rows)

    # Fetch MCQ questions based on the accepted student requests
    mcq_questions = []
     
    if not student_requests.empty:
        company_email = student_requests.iloc[0]['company_email']
        sql_mcq_questions = """
            SELECT * FROM mcqquestion 
            WHERE company_email=%s
        """
        mcq_questions = pd.read_sql_query(sql_mcq_questions, mydb, params=(company_email,))
        print(mcq_questions)
         
        # cursor = mydb.cursor(dictionary=True)
        # cursor.execute(sql_mcq_questions, (company_email,))
        # rows = cursor.fetchall()
        # mcq_questions = pd.DataFrame(rows)

    # Render template with MCQ questions
    return render_template(
        'viewmcqquestions.html', 
        cols=mcq_questions.columns.values, 
        rows=mcq_questions.values.tolist()
    )
    

@app.route('/stdview_mcq_questions')
def stdview_mcq_questions():
    sql = "SELECT * FROM mcqquestion WHERE company_email = %s"
    print(sql)
    company_email = session['compyemail']  # Temporarily hardcoded for testing purposes
    mycursor.execute(sql, (company_email,))
    data = mycursor.fetchall()
    print(data)  # Debugging line to check data fetched
    return render_template('stdview_mcq_questions.html', data=data)


@app.route("/submit_test", methods=["POST"])
def submit_test():
    # Extract answers from the form
    answers = request.form.to_dict()

    # Assuming you have the student email and company email from session or other source
    student_email = session['studentsemail']  # Replace with actual student email
    company_email = session['compyemail']  # Replace with actual company email

    for i in range(1, 6):  # Assuming there are 5 questions
        question_id = answers.get(f"question_id_{i}")
        code = answers.get(f"code_{i}")
        expected_output = answers.get(f"expected_output_{i}")
        output = ''  # You can add logic here to capture output if needed
        match = 'Yes' if output == expected_output else 'No'

        # Create a new coding result entry and add it to the database
        sql = """INSERT INTO coding_results (student_email, company_email, question_id, code, output, expected_output, `match`)
                         VALUES (%s, %s, %s, %s, %s, %s, %s)"""
        mycursor.execute(sql, (student_email, company_email, question_id, code, output, expected_output, match))
        mydb.commit()

    flash("Test submitted successfully!", "success")
    return render_template('studenthome.html', msg="Test submitted successfully!")


@app.route('/submit_answer', methods=['POST'])
def submit_answer():
    # Retrieve student details
    sql = "SELECT * FROM studentrequest WHERE student_email=%s"
    mycursor.execute(sql, (session['studentsemail'],))
    dc = mycursor.fetchall()
    if not dc:
        flash("Student not found!", "danger")
        return redirect(url_for('viewmcqquestions'))
    
    stdname = dc[0][1]
    stdemail = dc[0][2]
    company_email = dc[0][11]

    # Check if the student has already submitted answers for this company
    sql = "SELECT * FROM student_answers WHERE student_email=%s AND company_email=%s"
    mycursor.execute(sql, (stdemail, company_email))
    existing_answers = mycursor.fetchall()
    
    if existing_answers:
        flash("You have already submitted your answers for this company!", "danger")
        return redirect(url_for('viewmcqquestions'))

    # Retrieve questions for the company
    sql = "SELECT * FROM mcqquestion WHERE company_email=%s"
    mycursor.execute(sql, (company_email,))
    data = mycursor.fetchall()

    if not data:
        flash("No questions found for this company!", "danger")
        return redirect(url_for('viewmcqquestions'))

    for key in request.form:
        if key.startswith('question_'):
            question_id = key.split('_')[1]
            selected_answer = request.form.get(key)

            question = next((q for q in data if str(q[0]) == question_id), None)
            if question:
                question_text = question[2]
                correct_option = question[7]

                # Insert the answer into the database
                sql = """INSERT INTO student_answers (student_name, student_email, company_email, question_id, question_text, selected_answer, correct_answer)
                         VALUES (%s, %s, %s, %s, %s, %s, %s)"""
                mycursor.execute(sql, (stdname, stdemail, company_email, question_id, question_text, selected_answer, correct_option))
                mydb.commit()

    flash("Your answers have been submitted successfully!", "success")
    return redirect(url_for('viewmcqquestions'))



@app.route("/viewanswer")
def viewanswer():
    if 'company_name' not in session:
        return redirect(url_for('companysignin'))

    return render_template('viewanswer.html')

@app.route('/view_student_answers')
def view_student_answers():
    # Retrieve the student's answers from the database
    sql = "SELECT student_name, student_email, company_email, question_id, question_text, selected_answer, correct_answer FROM student_answers WHERE company_email=%s"
    mycursor.execute(sql, (session['compyemail'],))
    answers = mycursor.fetchall()

    # Calculate the score for each student
    student_scores = {}
    for answer in answers:
        student_email = answer[1]
        selected_answer = answer[5]
        correct_answer = answer[6]

        if student_email not in student_scores:
            student_scores[student_email] = {'total': 0, 'correct': 0}

        student_scores[student_email]['total'] += 1
        if selected_answer == correct_answer:
            student_scores[student_email]['correct'] += 1

    # Determine if each student qualifies for the coding test
    student_qualifications = {}
    for email, scores in student_scores.items():
        total_questions = scores['total']
        correct_answers = scores['correct']
        percentage = (correct_answers / total_questions) * 100
        session['student_percentage'] = percentage
        student_qualifications[email] = percentage >= 75
        print('percentage', percentage)
        print('yyyyyyyyy', student_qualifications[email])

    # Render the answers and qualifications in the template
    return render_template('viewmcqanswer.html', answers=answers, student_qualifications=student_qualifications)


@app.route('/next_exam/<student_email>')
def next_exam(student_email):
    student_email= session['studentsemail']
    sql = "update studentrequest set Status='nextexam' where student_email='%s'" % (student_email)
    mycursor.execute(sql)
    mydb.commit()
    flash("Proceed For Coding Test.", "success")
    return redirect(url_for('viewrequeststd'))


@app.route("/codingtest")
def codingtest():
    # Fetch drive data
    student_email = session['studentsemail']
    sql_drive = "SELECT * FROM studentrequest WHERE Status='nextexam' AND student_email=%s"
    data = pd.read_sql_query(sql_drive, mydb, params=(student_email,))
    print(data)

    # Fetch coding test data
    sql = "SELECT * FROM coding_tests WHERE company_email = %s"
    company_email = session['compyemail']
    mycursor.execute(sql, (company_email,))
    data1 = mycursor.fetchall()
    print(data1)  # Debugging line to check data fetched

    return render_template('codingtest.html', data1=data1)


def submit_coding_test():
    student_email = session.get('studentsemail')
    company_email = session.get('compyemail')
    language = request.form.get('testLanguage')
    results = []

    for i in range(1, 6):
        code = request.form.get(f'code_{i}')
        expected_output = request.form.get(f'expected_output_{i}')
        question_id = request.form.get(f'question_id_{i}')
        
        if code:
            output = execute_code(language, code)
            output_lines = output.strip().splitlines()
            expected_lines = expected_output.strip().splitlines()

            match = (output_lines == expected_lines)  # Update to strict equality for all lines

            result = {
                'student_email': student_email,
                'company_email': company_email,
                'question_id': question_id,
                'code': code,
                'output': output,
                'expected_output': expected_output,
                'match': match
            }
            results.append(result)
    
    # Save results to the database
    for result in results:
        try:
            sql = """
            INSERT INTO coding_results (student_email, company_email, question_id, code, output, expected_output, `match`)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
            """
            values = (result['student_email'], result['company_email'], result['question_id'], result['code'], result['output'], result['expected_output'], result['match'])
            mycursor.execute(sql, values)
            mydb.commit()
        except Exception as e:
            print(f"Database error: {str(e)}")

    flash('Your code has been submitted and evaluated.')
    return redirect(url_for('codingtest'))


import requests

def execute_code(language, code):
    url = "https://api.jdoodle.com/v1/execute"  # Make sure this is the correct endpoint
    client_id = "abc123xyz456"  # Replace with your actual client_id
    client_secret = "def789uvw012"  # Replace with your actual client_secret
    headers = {'Content-Type': 'application/json'}

    data = {
        'clientId': client_id,
        'clientSecret': client_secret,
        'script': code,
        'language': language,
        'versionIndex': '0'
    }

    try:
        response = requests.post(url, headers=headers, json=data)
        result = response.json()
        
        print("API Response:", result)  # Debugging line

        if response.status_code == 200:
            if 'output' in result:
                return result['output']
            else:
                return f"Error: {result.get('error', 'Unknown error')}"
        else:
            return f"API returned status code {response.status_code}: {result.get('error', 'Unknown error')}"
    except Exception as e:
        return f"Exception occurred: {str(e)}"


@app.route("/viewanswerforcode")
def viewanswerforcode():
    if 'company_name' not in session:
        return redirect(url_for('companysignin'))

    # Fetch the student's answers from the database
    sql = "SELECT student_email, question_id, code, output, expected_output, `match` FROM coding_results WHERE company_email=%s"
    mycursor.execute(sql, (session['compyemail'],))
    answers = mycursor.fetchall()

    return render_template('viewanswerforcode.html', answers=answers)

@app.route("/acceptcandidate/<id>")
def acceptcandidate(id):
    print(f"Candidate ID: {id}")
    
    # Fetch details from coding_results
    sql = "SELECT * FROM coding_results WHERE id = %s"
    mycursor.execute(sql, (id,))
    dc = mycursor.fetchall()
    
    if not dc:
        flash("No matching records found in coding results.", "error")
        return redirect(url_for('viewanswerforcode'))

    student_email = dc[0][1]
    print(f"Student Email: {student_email}")

    # Fetch company details
    sql = "SELECT * FROM company WHERE id = %s"
    mycursor.execute(sql, (id,))
    dc = mycursor.fetchall()
    
    if not dc:
        flash("No matching records found in company table.", "error")
        return redirect(url_for('viewanswerforcode'))

    company_name = dc[0][1]
    email = dc[0][2]
    phone = dc[0][3]
    print(f"Company Name: {company_name}, Company Email: {email}")

    # Fetch student request details
    sql = "SELECT * FROM studentrequest WHERE id = %s"
    mycursor.execute(sql, (id,))
    dc = mycursor.fetchall()
    
    if not dc:
        flash("No matching records found in student requests.", "error")
        return redirect(url_for('viewanswerforcode'))

    student_name = dc[0][1]
    position = dc[0][9]
    Status = "nextexam"
    print(f"Student Name: {student_name}, Position: {position}")

    # Prepare and send the email
    mail_content = f"""
    Dear {student_name},

    Congratulations!

    We are pleased to inform you that your application for the position of {position} at {company_name} has been selected. You have been chosen for the next steps in our recruitment process.

    Details:
    - Position: {position}
    - Company: {company_name}

    We will contact you soon with further details. If you have any questions, feel free to reach out.

    Best Regards,
    {company_name}
    Contact Information: {phone}
    Company Email: {email}
    """
    
    # Send email
    sender_address = 'cse.takeoff@gmail.com'
    sender_pass = 'digkagfgyxcjltup'
    receiver_address = student_email
    message = MIMEMultipart()
    message['From'] = sender_address
    message['To'] = receiver_address
    message['Subject'] = 'Campus Placement Recruitment System'
    message.attach(MIMEText(mail_content, 'plain'))
    session_email = smtplib.SMTP('smtp.gmail.com', 587)
    session_email.starttls()
    session_email.login(sender_address, sender_pass)
    text = message.as_string()
    session_email.sendmail(sender_address, receiver_address, text)
    session_email.quit()

    print(Status)
    # Update status in the database
    sql = "update studentrequest set Status='Selected' where id='%s'" % (id)
    mycursor.execute(sql)
    mydb.commit()

    flash("Drive request accepted successfully. Email sent to the student.", "success")
    return redirect(url_for('viewanswerforcode'))


@app.route("/rejectedcondidate/<id>")
def rejectedcondidate(id=0):
    print(f"Candidate ID: {id}")
    
    # Fetch details from coding_results
    sql = "SELECT * FROM coding_results WHERE id = %s"
    mycursor.execute(sql, (id,))
    dc = mycursor.fetchall()
    
    if not dc:
        flash("No matching records found in coding results.", "error")
        return redirect(url_for('viewanswerforcode'))

    student_email = dc[0][1]
    print(f"Student Email: {student_email}")

    # Fetch company details
    sql = "SELECT * FROM company WHERE id = %s"
    mycursor.execute(sql, (id,))
    dc = mycursor.fetchall()
    
    if not dc:
        flash("No matching records found in company table.", "error")
        return redirect(url_for('viewanswerforcode'))

    company_name = dc[0][1]
    email = dc[0][2]
    phone = dc[0][3]
    print(f"Company Name: {company_name}, Company Email: {email}")

    # Fetch student request details
    sql = "SELECT * FROM studentrequest WHERE id = %s"
    mycursor.execute(sql, (id,))
    dc = mycursor.fetchall()
    
    if not dc:
        flash("No matching records found in student requests.", "error")
        return redirect(url_for('viewanswerforcode'))

    student_name = dc[0][1]
    position = dc[0][9]
    print(f"Student Name: {student_name}, Position: {position}")

    mail_content = f"""
        Dear {student_name},

        Thank you for your interest in the {position} position at {company_name}.

        After careful consideration, we regret to inform you that we will not be moving forward with your application at this time.

        We appreciate your interest in our company and encourage you to apply for future openings.

        Best Wishes,
        {company_name}
        Contact Information: {phone}
        Company Email: {email}
        """
    sender_address = 'cse.takeoff@gmail.com'
    sender_pass = 'digkagfgyxcjltup'
    receiver_address = student_email
    message = MIMEMultipart()
    message['From'] = sender_address
    message['To'] = receiver_address
    message['Subject'] = 'Campus Placement Recruitment System'
    message.attach(MIMEText(mail_content, 'plain'))
    session_email = smtplib.SMTP('smtp.gmail.com', 587)
    session_email.starttls()
    session_email.login(sender_address, sender_pass)
    text = message.as_string()
    session_email.sendmail(sender_address, receiver_address, text)
    session_email.quit()


    sql = "update studentrequest set Status='NotSelected' where id='%s'" % (id)
    mycursor.execute(sql)
    mydb.commit()
    flash("Drive request rejected successfully.", "danger")
    return redirect(url_for('view_student_answers'))


@app.route("/viewselectedstudent")
def viewselectedstudent():
    # if 'student_email' not in session:
    #     return redirect(url_for('viewselectedstudent'))
    # Fetch drive data
    sql = " SELECT * FROM studentrequest WHERE student_email=%s AND Status='Selected'"
    data = pd.read_sql_query(sql, mydb, params=(session['studentsemail'],))

    return render_template('viewselectedstudent.html', cols=data.columns.values, rows=data.values.tolist())



##########autogenerate mcqs#########

GOOGLE_API_KEY = ' AIzaSyDPezGbu9hmt414dS5sHH4GP55ut9SWBKo'
# Assuming genai is configured properly
import google.generativeai as genai
genai.configure(api_key=GOOGLE_API_KEY)

# Initialize Gemini model
model = genai.GenerativeModel('gemini-2.0-flash')

def generate_mcqs(text, number, tone):
    prompt = f"""
    Text: {text}
    
    You are an expert MCQ maker. Given the above text, create a quiz of {number} multiple choice questions in {tone} tone.
    Format like the example:
    
    1: Question?
    a) Option 1
    b) Option 2
    c) Option 3
    d) Option 4
    Correct answer: x
    """
    
    for attempt in range(3):  # Retry up to 3 times
        try:
            response = model.generate_content(prompt)
            quiz = response.text
            
            # Parse generated quiz
            mcqs = quiz.split('\n\n')
            quiz_data = []
            answer_key = {}
            for mcq in mcqs:
                lines = mcq.split('\n')
                if len(lines) < 5:
                    continue
                question = lines[0].strip()
                options = {chr(ord('a') + i): line[3:].strip() for i, line in enumerate(lines[1:5])}
                correct = lines[5][14:].strip()
                quiz_data.append({"Question": question, "Options": options})
                answer_key[question] = correct
                
            return quiz_data, answer_key
        
        except Exception as e:
            if "500" in str(e):
                time.sleep(2 ** attempt)  # Exponential backoff
            else:
                return None, None
    
    return None, None
def parse_question_data(quiz_entry):
    # Extracting the question and options from the dictionary
    question = quiz_entry['Question']
    options = quiz_entry['Options']
    
    # Ensure options are correctly extracted, defaulting to empty string if missing
    opt1 = options.get('a', '')
    opt2 = options.get('b', '')
    opt3 = options.get('c', '')
    opt4 = options.get('d', '')
    
    return question, opt1, opt2, opt3, opt4

@app.route('/prediction', methods=['GET', 'POST'])
def prediction():
    if request.method == "POST":
        text = None
        file = request.files["file"]
        sub = request.form["sub"]
        txt = request.form["txt"]
        number_of_questions = int(request.form["number_of_questions"])
        tone = request.form["tone"]

        
        # File or URL or Text Area content retrieval
        if file:
            text = file.read().decode("utf-8")
        elif txt:
            text = txt

        # Generate MCQs
        quiz_data, answer_key = generate_mcqs(text, number_of_questions, tone)
        print(answer_key)
        if quiz_data:
            session['quiz_data'] = quiz_data
            session['answer_key'] = answer_key
            for i, quiz_entry in enumerate(quiz_data):
                parsed_data = parse_question_data(quiz_entry)
                
                if parsed_data:
                    question, opt1, opt2, opt3, opt4 = parsed_data
                    answer = answer_key[question]
                    answer = answer.split(' ')[-1]  # Extract the answer from the string
                    print(answer)
                    mycursor.execute('''
                    INSERT INTO mcqquestion (company_email, question_text, option1, option2, option3, option4, correct_option, stage)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, 'MCQ')
                    ''', (session['compyemail'], question, opt1, opt2, opt3, opt4, answer))
            mydb.commit()
            print("data stored successfully")
            return render_template("add_question2.html", quiz_data=quiz_data)
            

        else:
            flash("Failed to generate MCQs.")
            return redirect(url_for("index"))

    return render_template("add_question2.html")

@app.route("/download/<format>")
def download(format):
    quiz_data = session.get("quiz_data", [])
    answer_key = session.get("answer_key", {})

    if format == "pdf":
        pdf_file = SimpleDocTemplate('mcqs.pdf', pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()

        for mcq in quiz_data:
            elements.append(Paragraph(mcq['Question'], styles['Heading1']))
            for option, text in mcq['Options'].items():
                elements.append(Paragraph(f"{option}) {text}", styles['Normal']))
            elements.append(Paragraph("", styles['Normal']))

        pdf_file.build(elements)
        return send_file("mcqs.pdf", as_attachment=True)

    elif format == "word":
        document = Document()
        for mcq in quiz_data:
            document.add_heading(mcq['Question'], level=1)
            for option, text in mcq['Options'].items():
                document.add_paragraph(f"{option}) {text}")
            document.add_paragraph()

        document.save("mcqs.docx")
        return send_file("mcqs.docx", as_attachment=True)

    elif format == "answer_key":
        answer_key_doc = Document()
        answer_key_doc.add_heading("Answer Key", level=1)
        for question, answer in answer_key.items():
            answer_key_doc.add_paragraph(f"{question} - Correct answer: {answer}")

        answer_key_doc.save("answer_key.docx")
        return send_file("answer_key.docx", as_attachment=True)

    return redirect(url_for("results"))

@app.route('/autoGenerateCodingQuestions', methods=['GET', 'POST'])
def autoGenerateCodingQuestions():
    if request.method == "POST":
        Planguage = request.form["language"]
        number_of_questions = int(request.form["number_of_questions"])
        level = request.form["level"]

        # Generate Coding Questions
        quiz_data = code_q_generate(Planguage, number_of_questions, level)
        print('Generated Coding Questions:', quiz_data)

        if quiz_data:
            session['quiz_data'] = quiz_data

            for entry in quiz_data:
                question_text = entry.get("Question", "")
                sample_input = entry.get("Sample Input", "")
                sample_output = entry.get("Sample Output", "")
                print('question_text', question_text)
                print('sample_input', sample_input)
                print('sample_output', sample_output)

                # Insert into your coding question table
                mycursor.execute('''
                    INSERT INTO coding_tests (company_email, question_text, sample_input, expected_output,language)
                    VALUES (%s, %s, %s, %s, %s)
                ''', (session['compyemail'], question_text, sample_input, sample_output, Planguage))

            mydb.commit()
            print("Coding questions stored successfully")
            return render_template("add_coding_Q.html", quiz_data=quiz_data)

    return render_template("add_coding_Q.html")


#-------------------Auto generates coding function-----------------------------------
def code_q_generate(Planguage, number_of_questions, level):
    prompt = f"""
You are an expert coding question generator. Generate {number_of_questions} {Planguage} coding questions.
Each question must include:
1. Problem Statement
2. Sample Input
3. Sample Output

Format example:
Question: Write a Python program to check if a string is a palindrome.
Sample Input: "madam"
Sample Output: True
"""

    for attempt in range(3):
        try:
            response = model.generate_content(prompt)
            print("Raw AI Response:\n", response.text)

            lines = response.text.strip().split('\n')
            clean_lines = [line.replace("```", "").strip() for line in lines if line.strip()]

            quiz_data = []
            i = 0
            while i < len(clean_lines) - 2:
                # Check if proper keywords exist
                if ("input" in clean_lines[i + 1].lower() and 
                    "output" in clean_lines[i + 2].lower()):
                    
                    question = clean_lines[i].replace("Question:", "").strip()
                    sample_input = clean_lines[i + 1].split(":", 1)[-1].strip()
                    sample_output = clean_lines[i + 2].split(":", 1)[-1].strip()

                    quiz_data.append({
                        "Question": question,
                        "Sample Input": sample_input,
                        "Sample Output": sample_output
                    })
                    i += 3  # move to next question set
                else:
                    i += 1  # skip malformed lines

            return quiz_data

        except Exception as e:
            if "500" in str(e):
                time.sleep(2 ** attempt)
            else:
                print("Error:", e)
                return None

    return None

@app.route("/test_Responce")
def test_Responce():
    student_mail = session.get('student')
    student_test_per = session.get('student_percentage')

    if student_test_per is None:
        return "Student percentage not found in session.", 400

    # Determine the message
    if student_test_per >= 75:
        message = f"Congratulations! You scored {student_test_per:.2f}%. Great job!"
    else:
        improvement_needed = 75 - student_test_per
        message = (f"You scored {student_test_per:.2f}%. "
                   f"You need to improve your skills by {improvement_needed:.2f}% to reach 75%.")

    return render_template("test_Responce.html", 
                           student_mail=student_mail, 
                           percentage=student_test_per, 
                           message=message)

@app.route("/view_results")
def view_results():
    # Retrieve the student's answers from the database
    sql = "SELECT student_name, student_email, company_email, question_id, question_text, selected_answer, correct_answer FROM student_answers WHERE student_email=%s"
    mycursor.execute(sql, (session['student'],))
    answers = mycursor.fetchall()
    print(answers)
    print('sessions', session['student'])
    # Calculate the score for each student
    student_scores = {}
    for answer in answers:
        student_email = answer[1]
        print('dfdsdgf', student_email)
        selected_answer = answer[5]
        correct_answer = answer[6]

        if student_email not in student_scores:
            student_scores[student_email] = {'total': 0, 'correct': 0}

        student_scores[student_email]['total'] += 1
        if selected_answer == correct_answer:
            student_scores[student_email]['correct'] += 1

    # Determine if each student qualifies for the coding test
    student_qualifications = {}
    for email, scores in student_scores.items():
        total_questions = scores['total']
        correct_answers = scores['correct']
        percentage = (correct_answers / total_questions) * 100
        session['student_percentage'] = percentage
        student_qualifications[email] = percentage >= 75
        print('percentage', percentage)
        print('yyyyyyyyy', student_qualifications[email])

    # Render the answers and qualifications in the template
    return render_template('view_results.html', answers=answers, student_qualifications=student_qualifications)




if __name__ == "__main__":
    app.run(debug=True)
